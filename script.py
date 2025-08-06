from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import pandas as pd
from datetime import datetime

# --- Dicionário de Cores ---
cores_por_tema = {
    "CONHECIMENTO E INOVAÇÃO": "4400FF",  # Azul
    "SAÚDE E QUALIDADE DE VIDA": "ED282C",  # Vermelho
    "SEGURANÇA E CIDADANIA": "FFB000",  # Amarelo
    "DESENVOLVIMENTO SUSTENTÁVEL": "87D200",  # Verde
    "Gestão, Transparência e Participação": "002060"  # Azul escuro
}

# --- Funções Auxiliares ---
def set_cell_background(cell, hex_color):
    """Define a cor de fundo para uma célula da tabela."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace("#", ""))
    tcPr.append(shd)

def set_paragraph_background(paragraph, color):
    """Define a cor de fundo para um parágrafo."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color.replace("#", ""))
    pPr.append(shd)

# --- Leitura e Preparação dos Dados ---
# Carregar o arquivo Excel (certifique-se de que o caminho está correto)
df = pd.read_excel('Iniciativas - RGS 2025.1 - Extração Painel de Controle copy.xlsx', skiprows=1)

# Selecionar e renomear as colunas
colunas = ['Órgão', 'Iniciativa', 'Status Informado', 'Ação', 'Programa',
           'Início Realizado', 'Término Realizado', 'RGS 2025.1 - GGGE', 'Localização Geográfica', 'Objetivo Estratégico']
df2 = df[colunas]

df2.rename(columns={
    'Órgão': 'Orgao',
    'Iniciativa': 'Iniciativa',
    'Status Informado': 'Status_Informado',
    'Ação': 'Acao',
    'Programa': 'Programa',
    'Início Realizado': 'Inicio_Realizado',
    'Término Realizado': 'Termino_Realizado',
    'RGS 2025.1 - GGGE': 'RGS_2025_GGGE',
    'Localização Geográfica': 'Localizacao_Geografica',
    'Objetivo Estratégico': 'Objetivo_Estrategico'
}, inplace=True)

# Converter as colunas de datas
df2[['Inicio_Realizado', 'Termino_Realizado']] = df2[['Inicio_Realizado', 'Termino_Realizado']].apply(
    lambda x: pd.to_datetime(x, errors='coerce', dayfirst=True)
)

# --- Criação do Documento Word ---
doc = Document()

for idx, row in enumerate(df2.itertuples()):
    if idx > 0:
        doc.add_paragraph('\n')

    cor = cores_por_tema.get(row.Objetivo_Estrategico, "D3D3D3")

    p_orgao = doc.add_paragraph()
    p_orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_orgao.paragraph_format.space_before = Pt(0)
    p_orgao.paragraph_format.space_after = Pt(0)
    run = p_orgao.add_run(str(row.Orgao))
    run.font.name = 'Gilroy ExtraBold'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_paragraph_background(p_orgao, 'D3D3D3')

    p_programa = doc.add_paragraph()
    p_programa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_programa.paragraph_format.space_before = Pt(0)
    p_programa.paragraph_format.space_after = Pt(0)
    run = p_programa.add_run(str(row.Programa))
    run.font.name = 'Gilroy Light'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_paragraph_background(p_programa, cor)

    p_acao = doc.add_paragraph()
    p_acao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_acao.paragraph_format.space_before = Pt(0)
    p_acao.paragraph_format.space_after = Pt(0)
    run = p_acao.add_run(str(row.Acao))
    run.font.name = 'Gilroy Light'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_paragraph_background(p_acao, cor)

    doc.add_paragraph()

    status_imagem = 'imagens/concluído.png' if row.Status_Informado == 'CONCLUÍDO' else 'imagens/em_excecucao.png'
    status_texto_label = 'Data de Entrega:' if row.Status_Informado == 'CONCLUÍDO' else 'Data de Início:'
    prazo = row.Termino_Realizado if row.Status_Informado == 'CONCLUÍDO' else row.Inicio_Realizado
    icone_localizacao_path = 'imagens/localização.png'
    icone_calendario_path = 'imagens/calendário.png'

    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell_iniciativa = table.cell(0, 0).merge(table.cell(0, 4))
    p_iniciativa = cell_iniciativa.paragraphs[0]
    p_iniciativa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_iniciativa.add_run(str(row.Iniciativa))
    font_iniciativa = run.font
    font_iniciativa.name = 'Gilroy ExtraBold'
    font_iniciativa.size = Pt(10)
    font_iniciativa.bold = True
    font_iniciativa.color.rgb = RGBColor(0, 0, 0)
    cell_iniciativa.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_background(cell_iniciativa, 'D3D3D3')

    cell_status = table.cell(1, 0).merge(table.cell(1, 1))
    p_status = cell_status.paragraphs[0]
    run_label_status = p_status.add_run()
    run_label_status.add_picture(status_imagem, width=Inches(0.17))
    run_label_status.add_text('  Status:  ')
    font_label_status = run_label_status.font
    font_label_status.name = 'Neutro Thin'
    font_label_status.size = Pt(9)
    font_label_status.color.rgb = RGBColor(0, 0, 0)
    run_valor_status = p_status.add_run(str(row.Status_Informado))
    font_valor_status = run_valor_status.font
    font_valor_status.name = 'Neutro'
    font_valor_status.size = Pt(10)
    
    # --- SEÇÃO MODIFICADA PARA APROXIMAR A DATA ---
    # Agora mesclamos as colunas 2, 3 e 4 para colocar tudo junto.
    cell_data_merged = table.cell(1, 2).merge(table.cell(1, 4))
    p_data = cell_data_merged.paragraphs[0]

    # Adiciona o rótulo da data
    run_data_label = p_data.add_run()
    run_data_label.add_picture(icone_calendario_path, width=Inches(0.17))
    run_data_label.add_text(f'  {status_texto_label} ') # Adicionado um espaço no final
    font_data_label = run_data_label.font
    font_data_label.name = 'Neutro Thin'
    font_data_label.size = Pt(9)

    # Adiciona o valor da data no mesmo parágrafo
    data_texto = prazo.strftime('%d/%m/%Y') if pd.notnull(prazo) else ''
    run_data_valor = p_data.add_run(f'\t\t {data_texto}')
    font_data_valor = run_data_valor.font
    font_data_valor.name = 'Neutro'
    font_data_valor.size = Pt(10)
    # --- FIM DA SEÇÃO MODIFICADA ---

    cell_loc_label = table.cell(2, 0).merge(table.cell(2, 1))
    p_loc_label = cell_loc_label.paragraphs[0]
    run_loc_label = p_loc_label.add_run()
    run_loc_label.add_picture(icone_localizacao_path, width=Inches(0.17))
    run_loc_label.add_text('  Municípios Atendidos: ')
    font_loc_label = run_loc_label.font
    font_loc_label.name = 'Neutro Thin'
    font_loc_label.size = Pt(9)

    cell_loc_valor = table.cell(2, 2).merge(table.cell(2, 4))
    p_loc_valor = cell_loc_valor.paragraphs[0]
    localizacao_texto = "" if pd.isnull(row.Localizacao_Geografica) else str(row.Localizacao_Geografica)
    run_loc_valor = p_loc_valor.add_run(localizacao_texto)
    font_loc_valor = run_loc_valor.font
    font_loc_valor.name = 'Neutro'
    font_loc_valor.size = Pt(10)
    
    cell_rgs = table.cell(3, 0).merge(table.cell(3, 4))
    p_rgs = cell_rgs.paragraphs[0]
    run_rgs = p_rgs.add_run(str(row.RGS_2025_GGGE))
    font_rgs = run_rgs.font
    font_rgs.name = 'Neutro'
    font_rgs.size = Pt(9)


# Salvar o documento final
doc.save('teste.docx')


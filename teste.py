from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn  
import pandas as pd
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from datetime import datetime


cores_por_tema = {
    "CONHECIMENTO E INOVAÇÃO": "#4400FF",  # Azul
    "SAÚDE E QUALIDADE DE VIDA": "#ED282C",  # Vermelho
    "SEGURANÇA E CIDADANIA": "#FFB000",  # Amarelo
    "DESENVOLVIMENTO SUSTENTÁVEL": "#87D200",  # Verde
    "Gestão, Transparência e Participação": "#002060"  # Azul escuro
}



def set_paragraph_background(paragraph, color):
    """
    Define a cor de fundo (background) para um parágrafo.
    :param paragraph: Objeto do parágrafo (docx.paragraph.Paragraph)
    :param color: Código hexadecimal para a cor de fundo (ex.: 'FFFF00' para amarelo)
    """
    # Obtém o elemento XML subjacente do parágrafo
    p = paragraph._p
    pPr = p.get_or_add_pPr()  # Adiciona ou obtém as propriedades do parágrafo

    # Cria um elemento <w:shd> para aplicar a cor de fundo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')  # Define o preenchimento como "clear"
    shd.set(qn('w:color'), 'auto')  # Define a cor do texto como padrão (automática)
    shd.set(qn('w:fill'), color)  # Configura o preenchimento do fundo com a cor escolhida (hexadecimal)

    # Adiciona o elemento <w:shd> nas propriedades do parágrafo
    pPr.append(shd)

# Carregar o arquivo Excel
df = pd.read_excel('Iniciativas - RGS 2025.1 - Extração Painel de Controle.xlsx')

# Selecionar e renomear as colunas
colunas = ['Órgão', 'Iniciativa', 'Status Informado', 'Ação', 'Programa',
           'Início Realizado', 'Término Realizado', 'RGS 2025.1 - GGGE', 'Localização Geográfica','Objetivo Estratégico']
df2 = df[colunas]

df2.rename(columns={
    'Órgão': 'Orgao',
    'Iniciativa': 'Iniciativa',
    'Status Informado': 'Status_Informado',
    'Ação': 'Acao',
    'Programa': 'Programa',
    'Início Realizado': 'Inicio_Realizado',
    'Término Realizado': 'Termino_Realizado',
    'RGS 2025.1 - GGGE': 'RGS_2025_GGGE',
    'Localização Geográfica': 'Localizacao_Geografica',
    'Objetivo Estratégico':'Objetivo_Estrategico'
}, inplace=True)

# Converter as colunas de datas
df2.loc[:, ['Inicio_Realizado', 'Termino_Realizado']] = df2[['Inicio_Realizado', 'Termino_Realizado']].apply(
    lambda x: pd.to_datetime(x, errors='coerce', dayfirst=True)
)

# Criação do documento
doc = Document()


for idx, row in enumerate(df2.itertuples()):
    if idx > 0:
        doc.add_paragraph('\n')  # Espaço entre blocos

    cor = cores_por_tema[row.Objetivo_Estrategico]

    # --- ORGAO ---
    p_orgao = doc.add_paragraph()
    p_orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_orgao.paragraph_format.space_before = Pt(0)
    p_orgao.paragraph_format.space_after = Pt(0)
    run = p_orgao.add_run(f'{row.Orgao}')
    run.font.name = 'Gilroy ExtraBold'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    set_paragraph_background(p_orgao, 'D3D3D3')

    # --- PROGRAMA ---
    p_programa = doc.add_paragraph()
    p_programa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_programa.paragraph_format.space_before = Pt(0)
    p_programa.paragraph_format.space_after = Pt(0)
    run = p_programa.add_run(f'{row.Programa}')
    run.font.name = 'Gilroy Light'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_paragraph_background(p_programa, cor)

    # --- ACAO ---
    p_acao = doc.add_paragraph()
    p_acao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_acao.paragraph_format.space_before = Pt(0)
    p_acao.paragraph_format.space_after = Pt(0)
    run = p_acao.add_run(f'{row.Acao}')
    run.font.name = 'Gilroy Light'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_paragraph_background(p_acao, cor)

    # Espaço menor que uma linha entre Acao e Iniciativa
    doc.add_paragraph()

    #Criando tabela
    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 1 linha
    cell = table.cell(0, 0)
    cell_merge = cell.merge(table.cell(0, 4))
    paragraph = cell_merge.paragraphs[0]
    paragraph.alignment = 1  # Center
    run = paragraph.add_run(str(row.Iniciativa))
    run.font.name = 'Gilroy ExtraBold'
    run.font.size = Pt(10)
    run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 2 linha
    imagem = 'imagens\concluído.png' if row.Status_Informado == 'CONCLUÍDO' else 'imagens\em_excecucao.png'
    data = 'Data de Entrega' if  row.Status_Informado == 'CONCLUÍDO' else 'Data de Início'
    valor_da_data = row.Termino_Realizado if  row.Status_Informado == 'CONCLUÍDO' else row.Inicio_Realizado
    if pd.notnull(valor_da_data):
        valor_da_data_str = valor_da_data.strftime('%d/%m/%Y')
    else:
        valor_da_data_str = ""
    p_img = table.cell(1, 0).paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(imagem, width=Inches(0.2))
    table.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # (Coluna 2) Texto do status
    p_status = table.cell(1, 1).paragraphs[0]
    p_status.add_run(f"Status: {row.Status_Informado}").font.size = Pt(9)
    table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # (Coluna 3) Ícone data
    p_date_icon = table.cell(1, 2).paragraphs[0]
    p_date_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date_icon = p_date_icon.add_run()
    run_date_icon.add_picture('imagens\calendário.png', width=Inches(0.2))
    table.cell(1, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # (Coluna 4) Texto "Data da Entrega:"
    p_data_label = table.cell(1, 3).paragraphs[0]
    p_data_label.add_run(f"{data}:").font.size = Pt(9)
    table.cell(1, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # (Coluna 5) Valor da data
    p_data_value = table.cell(1, 4).paragraphs[0]
    p_data_value.add_run(f'{valor_da_data}').font.size = Pt(9)
    table.cell(1, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    

    # 4ª linha: Descrição (merge)
    cell_desc = table.cell(3, 0)
    cell_desc.merge(table.cell(3, 4))
    p_desc = cell_desc.paragraphs[0]
    p_desc.add_run(f'{row.RGS_2025_GGGE}').font.size = Pt(9)
    cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.TOP




# Salvar o documento
doc.save("teste2.docx")  